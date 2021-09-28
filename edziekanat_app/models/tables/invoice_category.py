import json
import re

from ckeditor.fields import RichTextField
from django.core.exceptions import ValidationError
from django.db.models import *
from django.utils.translation import gettext as _
from docx import Document
#from edziekanat_app.forms import bind

BASE_PATTERN = r'(?s)(?<={{).*?(?=}})'

def regex_result_to_dict(lst):
    result = dict()
    for i in range(len(lst)):
        field_type = match_field_type(lst[i][0:3].lower())
        result[f'{field_type}_field_{i}'] = lst[i]
    return result


def find_pattern(text, pattern): return re.findall(pattern, text)


def match_field_type(item: str):
    return {
        'tex': 'text',
        'are': 'areatext',
        'dat': 'date',
        'pho': 'text',
        'val': 'value',
        'rad': 'radio',
        'che': 'check',
        'sel': 'select',
        'res': 'result',
        'fil': 'file',
        'que': 'query'
    }[item]


class InvoiceCategory(Model):

    def validate_file_extension(value):
        if not value.name.endswith('.docx'):
            raise ValidationError(u'Plik musi mieć rozszerzenie .docx')

    name = CharField(max_length=100,
                     unique=True,
                     verbose_name=_('Nazwa'))

    decision_query = CharField(max_length=100,
                               default="NONE",
                               verbose_name=_('Query osoby do osoby decyzyjnej'))

    field = ForeignKey(to='edziekanat_app.InvoiceField',
                       verbose_name=_('Dziedzina'),
                       on_delete=CASCADE, default=None, null=None)

    faq_link = URLField(blank=True,
                        help_text="Link do regulaminu",
                        verbose_name=_('Regulamin'))

    description = RichTextField(verbose_name=_('Opis'),
                                blank=True)

    docx_template = FileField(upload_to="edziekanat_app/invoice_templates",
                              verbose_name=_('Plik wzorcowy'),
                              validators=[validate_file_extension])

    dynamic_fields = {}

    field_types = CharField(max_length=1000, default=json.dumps(dynamic_fields), editable=False)

    def set_field_types(self, x):
        self.field_types = json.dumps(x)

    def get_field_types(self):
        return json.loads(self.field_types)

    def __str__(self):
        return self.name

    def save(self, *args, **kwargs):
        super().save(*args, **kwargs)
        self.document_parse()

        super(InvoiceCategory, self).save()

    class Meta:
        db_table = "edziekanat_app_invoice_categories"
        verbose_name = "Kategoria wniosku"
        verbose_name_plural = "Kategorie wniosków"

    def document_parse(self):
        doc = Document(self.docx_template.path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        concatted_text = '\n'.join(text)

        matches = find_pattern(concatted_text, BASE_PATTERN)

        result = regex_result_to_dict(matches)
        result = bind(result)
        self.set_field_types(result)
        replace_document_tags(doc, result, self.docx_template.path).save(self.docx_template.path)

def bind(init):
    results = {}
    for k in init:
        results.setdefault(k, []).append(init[k])
    return results

def replace_document_tags(doc: Document, dictionary: dict, final: bool = True):
    for i in dictionary:
        if final:
            to_replace = '{{' + i + '}}'
            match = "{{" + dictionary[i][0] + "}}"
        else:
            to_replace = str(dictionary[i][0])
            match = '{{' + i + '}}'
        for p in doc.paragraphs:
            matches = p.text.find(match)
            if matches >= 0:
                p.text = p.text.replace(match, to_replace)
        for table in doc.tables:
            found = False
            for row in table.rows:
                if found: break
                for cell in row.cells:
                    matches = cell.text.find(match)
                    if matches >= 0:
                        cell.text = cell.text.replace(match, to_replace)
                        found = True

    return doc


