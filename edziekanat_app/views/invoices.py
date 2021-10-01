import datetime
import os

from django.contrib import messages
from django.core.files import File
from django.core.files.storage import FileSystemStorage
from django.http import JsonResponse, Http404, HttpResponse
from django.shortcuts import render, redirect
from docx import Document
from formtools.wizard.views import SessionWizardView

from edziekanat_app.forms import AddInvoiceCategory, FormFileUpload, AcceptInvoiceForm, RejectInvoiceForm
from edziekanat_app.models.tables.invoices.invoice import Invoice
from edziekanat_app.models.tables.invoices.invoice_category import InvoiceCategory
from edziekanat_app.models.tables.invoices.invoice_field import InvoiceField
from edziekanat_app.utils import replace_document_tags, bind, get_query


def invoices_list(request):
    users_invoices = Invoice.objects.filter(created_by=request.user)
    for invoice in users_invoices:
        invoice.seen = True
        invoice.save()
    return render(request, 'user/invoices_list.html', context={'invoices': users_invoices,
                                                               'form_file_upload': FormFileUpload(),
                                                               'form_reject': RejectInvoiceForm(),})


def manage_invoices(request):
    if request.method == 'POST':
        action = request.POST.get('action')
        if action is None:
            messages.error(request, "Wystąpił błąd podczas wysyłania żądania.")
        else:
            invoice = Invoice.objects.filter(id=request.POST.get('id')).first()
            if action == 'REJECT':
                invoice.status = "Odrzucony"
                invoice.decision = request.POST.get('decision')
            elif action == 'ACCEPT':
                invoice.status = "Zaakceptowany"

            invoice.decision_author = request.user
            invoice.decision_date = datetime.datetime.now()
            invoice.save()

    return render(request, 'employee/manage_invoices.html',
                  context={'invoices': Invoice.objects.all(),
                           'form_reject': RejectInvoiceForm(),
                           'form_accept': AcceptInvoiceForm(),
                           'form_file_upload': FormFileUpload()})


def invoices(request):
    fields = InvoiceField.objects.all()
    invoice_fields = {}
    for field in fields:
        field_invoices = InvoiceCategory.objects.filter(field=field)
        if field_invoices.count() > 0:
            invoice_fields[field] = field_invoices

    return render(request, 'user/invoices.html', context={'invoices': invoice_fields})


def invoice_download(request):
    if request.method == 'GET' and request.is_ajax():
        id = request.GET.get('id', None)
        invoice = Invoice.objects.filter(id=id).first()
        file_path = invoice.invoice_file.path
        if os.path.exists(file_path):
            with open(file_path, 'rb') as fh:
                response = HttpResponse(fh.read(),
                                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = 'attachment; filename=' + file_path
                return response
        raise Http404

    return JsonResponse({}, status=400)


def create_invoice(request):
    return render(request, 'employee/create_invoice_category.html')


def create_invoice_category(request):
    form = AddInvoiceCategory()

    if request.method == "POST":
        name = request.POST.get('name')
        decision_query = request.POST.get('decision_query')
        faq_link = request.POST.get('faq_link')
        description = request.POST.get('description')
        docx_template = request.FILES['docx_template']
        field = InvoiceField.objects.filter(id=int(request.POST.get('field'))).first()
        try:
            InvoiceCategory.objects.create(
                name=name,
                decision_query=decision_query,
                faq_link=faq_link,
                description=description,
                docx_template=docx_template
            ).save()
            messages.success(request, 'Wysłano wiadomość')
        except Exception as e:
            messages.warning(request, f'Bład: {e.args[0]}')

    return render(request, 'employee/create_invoice_category.html', context={'form': form})


class InvoiceCreator(SessionWizardView):
    template_name = "user/create_invoice.html"
    file_storage = FileSystemStorage(location='invoices/')

    def get_form_kwargs(self, step=None):
        self.extra_context = {'title': 'Wybierz dziedzinę wniosku'}
        kwargs = {}
        if step == '1':
            field = self.get_cleaned_data_for_step('0')['field']
            queryset = InvoiceCategory.objects.filter(field=field)
            kwargs.update({'categories': queryset, })
            self.extra_context.update({'title': "Wybierz kategorię wniosku"})
        if step == '2':
            category = self.get_cleaned_data_for_step('1')['category']
            kwargs.update({'category': category, 'user': self.request.user})
            self.extra_context.update({'title': "Uzupełnij wymagane informacje", 'category_name': category.name})
        return kwargs

    def done(self, form_list, **kwargs):
        category = self.get_cleaned_data_for_step('1')['category']

        result = form_list[2].cleaned_data

        last_result_field = None
        for key in result:
            if key.startswith('select'):
                last_result_field = result[key]
                result[key] = result[key].name
            if key.startswith('file'):
                result[key] = result[key].name
            if key.startswith('result'):
                result[key] = get_query(result[key], self.request.user, base=last_result_field)

        results = bind(result)

        inv = Invoice.objects.create(category=category,
                                     created_by=self.request.user,
                                     created_date=datetime.datetime.now(),
                                     decision_author=self.request.user)  # todo
        try:
            doc = Document(category.docx_template.path)
            new_invoice_file_name = f"{category.name.replace(' ', '_')}_ID_{inv.id}.docx"
            new_invoice_file_path = f"edziekanat_app/invoices/{new_invoice_file_name}"
            replace_document_tags(doc, results, False).save(new_invoice_file_path)
            file = File(open(new_invoice_file_path, 'rb'))
            inv.invoice_file.save(name=new_invoice_file_name, content=file)
            file.close()
            os.remove(new_invoice_file_path)
        except Exception as e:
            inv.delete()
            messages.error(self.request, f"Wystąpił błąd podczas tworzenia wniosku.")
            raise e
        messages.success(self.request, f"Pomyślnie utworzono {category.name.lower()}")
        return redirect('edziekanat_app:home')

def get_reject_info(request):
    if request.method == 'GET' and request.is_ajax():
        invoice_id = request.GET.get('id', None)
        if Invoice.objects.filter(id=invoice_id).exists():
            invoice = Invoice.objects.filter(id=invoice_id).first()
            response_object = {
                'category_name': invoice.category.name,
                'decision_author': invoice.decision_author.__str__(),
                'decision': invoice.decision,
                'decision_date': invoice.decision_date.strftime("%d.%m.%Y %H:%M:%S")
            }
            response = JsonResponse({'Valid': True,
                                     'data': response_object}, status=200)
            return response
        else:
            return JsonResponse({'Valid': False}, status=200)

    return JsonResponse({}, status=400)
