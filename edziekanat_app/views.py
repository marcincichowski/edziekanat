import datetime
import os
from django.db.models import Count
from django.db.models import Q
from django.contrib import messages
from django.contrib.auth import authenticate, login
from django.contrib.auth.views import auth_logout
from django.core.files import File
from django.core.files.storage import FileSystemStorage
from django.core.serializers import serialize
from django.db.models import Q
from django.http import JsonResponse
from django.shortcuts import render, redirect
from docx import Document
from formtools.wizard.views import SessionWizardView

from edziekanat_app.forms import *
from .forms import LoginForm, EditUserForm, AddInvoiceCategory
from .models.tables.invoice import Invoice
from .models.tables.invoice_category import InvoiceCategory
from .models.crud.reject_invoice_info import RejectInvoiceInfo
from .forms import CreateMessage

from .models.tables.invoice_category import replace_document_tags
from .models.tables.messages.message import Message
from .models.tables.users.employee import Employee
from .models.tables.users.student import Student
from .models.tables.users.user import User


def format_timedelta(delta):
    days, rest = divmod(delta.total_seconds(), 86400)
    hours, rest = divmod(rest, 3600)
    minutes, seconds = divmod(rest, 60)
    return '{:02} dni {:02}h {:02}min'.format(int(days), int(hours), int(minutes))

def get_average_decision_time():
    invoices = Invoice.objects.all()
    if invoices.count() == 0:
        return 'brak zakończonych wniosków'
    sum_time = datetime.timedelta(days=3, hours=2, minutes=21)
    counter = 0
    for invoice in invoices:
        if invoice.decision_date is None:
            continue
        start = invoice.created_date.replace(tzinfo=None)
        end = invoice.decision_date.replace(tzinfo=None)
        delta = end - start
        sum_time += end - start
        counter += 1
    if counter == 0:
        return '∞'

    return format_timedelta(sum_time/counter)


def get_top_invoices():
    invoices = Invoice.objects.all().values('category').annotate(total=Count('category')).order_by('total')
    top_five = []
    for invoice in invoices[:5]:
        top_five.append(InvoiceCategory.objects.filter(id=invoice['category']).first().name)
    return top_five


def index(request, *args, **kwargs):

    average_time = get_average_decision_time()
    top_invoices = get_top_invoices()

    context = {
        'open_invoices': get_open_invoices(request.user),
        'closed_invoices': get_decision_invoices(request.user),
        'new_mesages': Message.objects.filter(reciever=request.user, seen=False),
        'all_invoices': get_user_invoices(request.user),
        'message': Message.objects.all(),
        'average_decision_time': average_time,
        'top_invoices': top_invoices
    }
    return render(request, 'index.html', context=context)


def invoices(request, *args, **kwargs):
    fields = InvoiceField.objects.all()
    invoices = {}
    for field in fields:
        field_invoices = InvoiceCategory.objects.filter(field=field)
        if field_invoices.count() > 0:
            invoices[field] = field_invoices

    return render(request, 'user/invoices.html', context={'invoices': invoices})


def invoice_download(request, *args, **kwargs):
    if request.method == 'GET' and request.is_ajax():
        id = request.GET.get('id', None)
        if Invoice.objects.filter(id=id).exists():
            return JsonResponse({'Valid': True, 'data': serialize('json', Invoice.objects.filter(id=id))}, status=200)
        else:
            return JsonResponse({'Valid': False}, status=200)

    return JsonResponse({}, status=400)


def message_details(request, *args, **kwargs):
    if request.method == 'GET' and request.is_ajax():
        id = request.GET.get('id', None)
        if Message.objects.filter(id=id).exists():
            return JsonResponse({'Valid': True, 'data': serialize('json', Message.objects.filter(id=id))}, status=200)
        else:
            return JsonResponse({'Valid': False}, status=200)

    return JsonResponse({}, status=400)


def get_reject_info(request, *args, **kwargs):
    if request.method == 'GET' and request.is_ajax():
        id = request.GET.get('id', None)
        if Invoice.objects.filter(id=id).exists():
            invoice = Invoice.objects.filter(id=id).first()
            response_object = {
                'category_name': invoice.category.name,
                'decision_author': invoice.decision_author.__str__(),
                'decision': invoice.decision,
                'decision_date': invoice.decision_date.strftime("%d.%m.%Y %H:%M:%S")
            }
            response = JsonResponse({'Valid': True,
                                     'data': response_object}, status=200)
            return response
        else:
            return JsonResponse({'Valid': False}, status=200)

    return JsonResponse({}, status=400)


def invoices_list(request, *args, **kwargs):
    invoices = Invoice.objects.filter(created_by=request.user)
    return render(request, 'user/invoices_list.html', context={'invoices': invoices})


def manage_invoices(request, *args, **kwargs):
    if request.method == 'POST':
        action = request.POST.get('action')
        if action is None:
            messages.error(request, "Wystąpił błąd podczas wysyłania żądania.")
        else:
            invoice = Invoice.objects.filter(id=request.POST.get('id')).first()
            if action == 'REJECT':
                invoice.status = "Odrzucony"
                invoice.decision = request.POST.get('decision')
            elif action == 'ACCEPT':
                invoice.status = "Zaakceptowany"

            invoice.decision_author = request.user
            invoice.decision_date = datetime.datetime.now()
            invoice.save()

    return render(request, 'employer/manage_invoices.html',
                  context={'invoices': Invoice.objects.all(),
                           'form_reject': RejectInvoiceForm(),
                           'form_accept': AcceptInvoiceForm(),
                           'form_file_upload': FormFileUpload()})


def administrators(request, *args, **kwargs):
    users = User.objects.all()
    edit_form = EditUserForm()
    return render(request, 'admin/administrators.html', context={'users': users, 'form': edit_form})


def account(request, *args, **kwargs):
    return render(request, 'user/account.html')


def config(request, *args, **kwargs):
    form = SystemTools()
    if request.method == "POST":
        title = request.POST.get('title')
        text = request.POST.get('broadcast')
        if title is not None and text is not None:
            users = User.objects.filter(~Q(id=request.user.id))
            for user in users:
                Message.objects.create(
                    sender=request.user,
                    reciever=user,
                    seen=False,
                    message_text=text,
                    message_title=title,
                    created_date=datetime.datetime.now()
                ).save()
            messages.success(request, f"Wysłano wiadomość do {users.count()} użytkowników.")

    return render(request, 'admin/tools.html', context={'form': form, 'toolname': 'Wiadomość broadcast'})


def settings(request, *args, **kwargs):
    return render(request, 'user/settings.html')


def create_invoice(request, *args, **kwargs):
    return render(request, 'user/create_invoice.html')


def database(request, *args, **kwargs):
    if request.method == 'GET':
        form_chair = AddChair()
        form_department = AddDepartment()
        form_faculty = AddFaculty()
        form_course = AddCourse()
        form_employee = AddEmployee()
        form_invoice_categorie = AddInvoiceCategory()
        form_invoice_field = AddInvoiceField()
        form_invoice = AddInvoice()
        form_job = AddJob()
        form_message = AddMessage()
        form_role = AddRole()
        form_spectialization = AddSpectialization()
        form_student = AddStudent()
        form_study_mode = AddStudyMode()
        form_subject = AddSubject()
        return render(request, 'admin/database.html',
                      {'form_chair': form_chair, 'form_department': form_department, 'form_faculty': form_faculty,
                       'form_course': form_course, 'form_employee': form_employee,
                       'form_invoice_categorie': form_invoice_categorie,
                       'form_invoice_field': form_invoice_field, 'form_invoice': form_invoice, 'form_job': form_job,
                       'form_message': form_message, 'form_role': form_role,
                       'form_spectialization': form_spectialization,
                       'form_student': form_student, 'form_study_mode': form_study_mode, 'form_subject': form_subject})
    elif request.method == 'POST':
        raise Exception("Not implemented")
    return render(request, 'admin/database.html')


def user_login(request):
    template = 'auth/login.html'

    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            email = request.POST['email']
            password = request.POST['password']
            user = authenticate(email=email, password=password)
            print(user)
            if user is not None:
                login(request, user)
                print(f"Zalogowano:\n"
                      f"Email:{user.email}\n"
                      f"Hasło:{user.password}\n"
                      f"Imie:{user.first_name}\n"
                      f"Nazwisko:{user.last_name}")
                messages.success(request, "Pomyślnie zalogowano!")

                return redirect('/')
            else:
                messages.error(request, 'Podano nieprawidłowe dane logowania.')
                return render(request, 'auth/login.html', {
                    'form': form,
                })
    else:
        form = LoginForm()

    return render(request, template, {'form': form, "time": datetime.datetime.now()})


def user_logout(request):
    auth_logout(request)
    messages.success(request, "Pomyślnie wylogowano.")
    return redirect('edziekanat_app:home')


class InvoiceCreator(SessionWizardView):
    template_name = "user/create_invoice.html"
    file_storage = FileSystemStorage(location='invoices/')

    def get_form_kwargs(self, step=None):
        self.extra_context = {'title': 'Wybierz dziedzinę wniosku'}
        kwargs = {}
        if step == '1':
            field = self.get_cleaned_data_for_step('0')['field']
            queryset = InvoiceCategory.objects.filter(field=field)
            kwargs.update({'categories': queryset, })
            self.extra_context.update({'title': "Wybierz kategorię wniosku"})
        if step == '2':
            category = self.get_cleaned_data_for_step('1')['category']
            kwargs.update({'category': category, 'user': self.request.user})
            self.extra_context.update({'title': "Wybrany wniosek", 'category_name': category.name})
        return kwargs

    def done(self, form_list, **kwargs):
        category = self.get_cleaned_data_for_step('1')['category']

        result = form_list[2].cleaned_data
        attachement = None  # not handling attachements yet

        last_result_field = None
        for key in result:
            if key.startswith('select'):
                last_result_field = result[key]
                result[key] = result[key].name
            if key.startswith('file'):
                attachement = result[key]
                result[key] = result[key].name
            if key.startswith('result'):
                result[key] = get_query(result[key], self.request.user, base=last_result_field)

        results = bind(result)

        inv = Invoice.objects.create(category=category,
                                     created_by=self.request.user,
                                     created_date=datetime.datetime.now(),
                                     decision_author=self.request.user)  # todo
        try:
            doc = Document(category.docx_template.path)
            new_invoice_file_name = f"{category.name.replace(' ', '_')}_ID_{inv.id}.docx"
            new_invoice_file_path = f"edziekanat_app/invoices/{new_invoice_file_name}"
            replace_document_tags(doc, results, False).save(new_invoice_file_path)
            file = File(open(new_invoice_file_path, 'rb'))
            inv.invoice_file.save(name=new_invoice_file_name, content=file)
            file.close()
            os.remove(new_invoice_file_path)
        except Exception as e:
            inv.delete()
            messages.error(self.request, f"Wystąpił błąd podczas tworzenia wniosku.")
            raise e
        messages.success(self.request, f"Pomyślnie utworzono {category.name.lower()}")
        return redirect('edziekanat_app:home')


def read_messages(request):
    messages = Message.objects.filter(reciever=request.user)
    for mess in messages:
        mess.seen = True
        mess.save()
    return redirect('/inbox')


def inbox(request):
    if request.method == 'POST':
        title = request.POST.get('message_title')
        text = request.POST.get('message_text')
        reciever = request.POST.get('reciever')
        if title is not None:
            if reciever == str(request.user.id):
                messages.warning(request, 'Nie możesz wysłać wiadomości do siebie.')
            else:
                try:
                    Message.objects.create(message_text =text,
                                           message_title = title,
                                           created_date = datetime.datetime.now(),
                                           reciever_id = reciever,
                                           sender_id = request.user.id).save()
                    messages.success(request, 'Wysłano wiadomość')
                except Exception as e:
                    messages.warning(request, f'Bład: {e.args[0]}')

    inbox_messages = Message.objects.filter(reciever=request.user).order_by('-created_date')
    return render(request, 'user/inbox.html', context={'inbox_messages': inbox_messages,
                                                       'new_message_form': CreateMessage()})


class UserCreator(SessionWizardView):
    template_name = "auth/register.html"

    def get_form_kwargs(self, step=None):
        kwargs = {}
        if step == '1':
            role = self.get_cleaned_data_for_step('0')['role']
            kwargs.update({'role': role, })
        return kwargs

    def done(self, form_list, **kwargs):
        email = self.get_cleaned_data_for_step('0')['email']
        password = self.get_cleaned_data_for_step('0')['password']
        password_repeat = self.get_cleaned_data_for_step('0')['password_repeat']
        role = self.get_cleaned_data_for_step('0')['role']
        first_name = self.get_cleaned_data_for_step('0')['first_name']
        last_name = self.get_cleaned_data_for_step('0')['last_name']
        address = self.get_cleaned_data_for_step('1')['address']
        phone = self.get_cleaned_data_for_step('1')['phone']
        birth_date = self.get_cleaned_data_for_step('1')['birth_date']
        allow_email_send = self.get_cleaned_data_for_step('1')['allow_email_send']

        if User.objects.filter(email=email).exists():
            messages.error(self.request, 'Użytkownik o takim adresie e-mail już istnieje.')
            return redirect('edziekanat_app:user_register')
        elif password != password_repeat:
            messages.error(self.request, 'Hasła nie zgadzają się.')
            return redirect('edziekanat_app:user_register')
        else:
            extra = {}
            if role.name == Student.base_role:
                extra = {
                    'course': self.get_cleaned_data_for_step('1')['course'],
                    'specialization': self.get_cleaned_data_for_step('1')['specialization'],
                }
            if role.name == Employee.base_role:
                extra = {
                    'job': self.get_cleaned_data_for_step('1')['job']
                }
            user = User.objects.create_user(
                password=password,
                email=email,
                role=role,
                phone=phone,
                address=address,
                birth_date=birth_date,
                allow_email_send=allow_email_send,
                extra=extra
            )
            user.first_name = first_name
            user.last_name = last_name
            user.save()
            print(f"Zarejestrowano:\n"
                  f"Email:{user.email}\n"
                  f"Hasło:{user.password}\n"
                  f"Imie:{user.first_name}\n"
                  f"Nazwisko:{user.last_name}")
            login(self.request, user)
            messages.success(self.request, 'Pomyślnie zarejestrowano!')
        return redirect('edziekanat_app:home')


def session_context_processor(request):
    return {
        "time": datetime.datetime.now()
    }


def get_open_invoices(user: User): return Invoice.objects.filter(status="W trakcie", created_by=user)


def get_closed_invoices(user: User): return Invoice.objects.filter(status="Zamknięte", created_by=user)


def get_decision_invoices(user: User): return Invoice.objects.filter(Q(status="Odrzucony") | Q(status="Zaakceptowany"),
                                                                     created_by=user)


def get_user_invoices(user: User): return Invoice.objects.filter(created_by=user)


def create_invoice_category(request):
    form = AddInvoiceCategory()

    if request.method == "POST":
        name = request.POST.get('name')
        decision_query = request.POST.get('decision_query')
        faq_link = request.POST.get('faq_link')
        description = request.POST.get('description')
        docx_template = request.FILES['docx_template']
        field = InvoiceField.objects.filter(id=int(request.POST.get('field'))).first()
        try:
            InvoiceCategory.objects.create(
                name=name,
                decision_query=decision_query,
                faq_link=faq_link,
                description=description,
                docx_template=docx_template
            ).save()
            messages.success(request, 'Wysłano wiadomość')
        except Exception as e:
            messages.warning(request, f'Bład: {e.args[0]}')

    return render(request, 'user/create_invoice_category.html', context={'form': form})
