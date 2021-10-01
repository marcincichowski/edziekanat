import datetime

from docx import Document

from .models.tables.users.student import Student
from .models.tables.users.user import User


def bind(init):
    results = {}
    for k in init:
        results.setdefault(k, []).append(init[k])
    return results


def get_query(queries, user: User, base=None):
    single_queries = queries.split('.')
    result = []

    if base is None:
        if single_queries[0] == 'user':
            result.append(user)
        elif single_queries[0] == 'student':
            result.append(Student.objects.filter(user=user).first())
        elif single_queries[0] == 'system':
            if single_queries[1] == 'today':
                return datetime.datetime.today().strftime('%d.%m.%Y') + " r."
            if single_queries[1] == 'study_session':
                start = (3, 1)
                end = (9, 3)
                if start < (datetime.date.today().month, datetime.date.today().day) < end:
                    return "letniej"
                else:
                    return "zimowej"
        else:
            raise Exception(f"Invalid base_object request: {single_queries[0]}")
    else:
        result.append(base)

    single_queries = single_queries[1:]
    for request, i in zip(single_queries, range(len(single_queries))):
        base_object = result[i]
        response = getattr(base_object, request)
        if response is None:
            raise Exception(f"Unable to get response. Object: {base_object} Query: {request}")
        result.append(response)
    return result[-1]


def replace_document_tags(doc: Document, dictionary: dict, final: bool = True):
    for i in dictionary:
        if final:
            to_replace = '{{' + i + '}}'
            match = "{{" + dictionary[i][0] + "}}"
        else:
            to_replace = str(dictionary[i][0])
            match = '{{' + i + '}}'
        for p in doc.paragraphs:
            matches = p.text.find(match)
            if matches >= 0:
                p.text = p.text.replace(match, to_replace)
        for table in doc.tables:
            found = False
            for row in table.rows:
                if found: break
                for cell in row.cells:
                    matches = cell.text.find(match)
                    if matches >= 0:
                        cell.text = cell.text.replace(match, to_replace)
                        found = True
    return doc
